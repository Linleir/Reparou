from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def gerar_docx_reparou():
    doc = Document()
    
    # Cabeçalho
    t = doc.add_heading('Projeto Reparou', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Seção 1: Introdução
    doc.add_heading('1. Nome do Sistema e Propósito', level=1)
    doc.add_paragraph('O Reparou conecta clientes a assistências técnicas de informática, '
                      'focando em transparência e previsibilidade de custos.')
    
    # Seção 2: Proposta de Valor
    doc.add_heading('2. Proposta de Valor', level=1)
    valores = [
        'Reduzir assimetria de informação',
        'Previsibilidade financeira',
        'Incentivo à economia circular (manutenção vs descarte)'
    ]
    for v in valores:
        doc.add_paragraph(v, style='List Bullet')
        
    # Seção 3: Casos de Uso
    doc.add_heading('3. Casos de Uso', level=1)
    doc.add_paragraph('O sistema contempla 18 casos de uso, organizados em:')
    doc.add_paragraph('CRUD de Lojas, Serviços e Estimativas; Sistema de Favoritos e Chat.', style='List Bullet')

    doc.save('Projeto_Reparou_Organizado.docx')

gerar_docx_reparou()